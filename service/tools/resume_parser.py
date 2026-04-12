"""
service/tools/resume_parser.py
简历解析工具

支持格式：PDF, DOC, DOCX
提取简历文本内容供 AI 分析使用
"""

import os
import re
from pathlib import Path
from typing import Optional


class ResumeParser:
    """简历解析器"""

    # 支持的文件扩展名
    SUPPORTED_EXTENSIONS = {".pdf", ".doc", ".docx"}

    @classmethod
    def extract_text(cls, file_path: str) -> str:
        """
        从简历文件中提取文本

        Args:
            file_path: 简历文件路径

        Returns:
            提取的文本内容

        Raises:
            ValueError: 不支持的文件格式
            FileNotFoundError: 文件不存在
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"简历文件不存在：{file_path}")

        file_ext = Path(file_path).suffix.lower()

        if file_ext not in cls.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"不支持的文件格式：{file_ext}，请上传 PDF、DOC 或 DOCX 文件"
            )

        if file_ext == ".pdf":
            return cls._extract_pdf(file_path)
        elif file_ext in {".doc", ".docx"}:
            return cls._extract_docx(file_path)

        return ""

    @staticmethod
    def _extract_pdf(file_path: str) -> str:
        """提取 PDF 文本"""
        try:
            from PyPDF2 import PdfReader

            reader = PdfReader(file_path)
            text_parts = []

            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)

            return "\n".join(text_parts)
        except ImportError:
            raise ImportError("请安装 PyPDF2：pip install PyPDF2>=3.0.0")
        except Exception as e:
            raise RuntimeError(f"PDF 解析失败：{e}")

    @staticmethod
    def _extract_docx(file_path: str) -> str:
        """提取 DOCX 文本（支持 .docx 和 .doc 格式）"""
        try:
            from docx import Document

            try:
                doc = Document(file_path)
            except Exception as e:
                # .doc 格式（旧版 Word 二进制）尝试转换
                if file_path.lower().endswith(".doc"):
                    return cls._extract_old_doc(file_path)
                raise RuntimeError(f"无法读取 Word 文件：{e}")

            text_parts = []

            # 提取段落
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # 提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))

            return "\n".join(text_parts)
        except ImportError:
            raise ImportError("请安装 python-docx：pip install python-docx>=1.1.0")
        except Exception as e:
            # .doc 格式（旧版 Word 二进制）尝试转换
            if file_path.lower().endswith(".doc"):
                return cls._extract_old_doc(file_path)
            raise RuntimeError(f"DOCX 解析失败：{e}")

    @classmethod
    def _extract_old_doc(cls, file_path: str) -> str:
        """提取旧版 .doc 格式（仅 Windows）"""
        if os.name != "nt":
            raise RuntimeError(
                "旧版 .doc 格式仅在 Windows 上支持，请转换为 .docx 或 .pdf"
            )

        try:
            import win32com.client
            import pythoncom

            pythoncom.CoInitialize()

            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False

            try:
                doc = word.Documents.Open(os.path.abspath(file_path))
                text = doc.Content.Text
                doc.Close(False)
                return text.strip()
            finally:
                word.Quit()
        except ImportError:
            raise RuntimeError(
                "请安装 pywin32 支持 .doc 格式：pip install pywin32\n"
                "或直接将简历转换为 .docx 或 .pdf 格式"
            )
        except Exception as e:
            raise RuntimeError(f".doc 读取失败：{e}")

    @classmethod
    def clean_text(cls, text: str) -> str:
        """
        清理简历文本

        - 移除多余空白
        - 移除特殊字符
        - 限制长度（避免超出 token 限制）
        """
        # 移除多余空白
        text = re.sub(r"\s+", " ", text)

        # 移除特殊字符，保留基本标点
        text = re.sub(
            r"[^\u4e00-\u9fa5a-zA-Z0-9\s\n\.\,\;\:\(\)\[\]\{\}\-\/\\@\#\$\%\&\*\+\=\?\!]",
            "",
            text,
        )

        # 限制长度（约 8000 字符，可根据模型调整）
        max_length = 8000
        if len(text) > max_length:
            text = text[:max_length] + "\n\n[简历内容过长，已截断]"

        return text.strip()

    @classmethod
    def parse_resume(cls, file_path: str) -> dict:
        """
        完整解析简历

        Returns:
            {
                "raw_text": 原始文本,
                "cleaned_text": 清理后的文本,
                "file_name": 文件名,
                "file_size": 文件大小(字节),
                "char_count": 字符数
            }
        """
        raw_text = cls.extract_text(file_path)
        cleaned_text = cls.clean_text(raw_text)

        return {
            "raw_text": raw_text,
            "cleaned_text": cleaned_text,
            "file_name": os.path.basename(file_path),
            "file_size": os.path.getsize(file_path),
            "char_count": len(cleaned_text),
        }
